import logging
import streamlit as st
import difflib
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pdf2image import convert_from_path
import pytesseract
import openai
import spacy
import os
import cv2
import numpy as np
import pandas as pd

# Set up logging for error handling
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


# Define error handling function
def handle_error(error_message):
    logging.error(error_message)
    st.error(error_message)


# Load NLP model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    import sys

    subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

# Set OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")


def extract_entities(text):
    """
    Extract named entities from a document text using spaCy.
    """
    doc = nlp(text)
    return [ent.text for ent in doc.ents]


def export_to_word(aligned_diff):
    doc = Document()
    doc.add_heading('Document Comparison Report', 0)

    # Add a section for the text comparison results
    doc.add_heading('Text Comparison', level=1)

    for old_line, new_line, tag in aligned_diff:
        if tag == "equal":
            doc.add_paragraph(f"Old: {old_line}")
            doc.add_paragraph(f"New: {new_line}")
        elif tag == "replace":
            doc.add_paragraph(f"Old (replaced): {old_line}", style='Color')
            doc.add_paragraph(f"New (replaced): {new_line}", style='Color')

    # Add a section for the table comparison results
    doc.add_heading('Table Comparison', level=1)
    for table_diff in aligned_diff:
        for old_cell, new_cell, tag in table_diff:
            if tag == "equal":
                doc.add_paragraph(f"Old: {old_cell}")
                doc.add_paragraph(f"New: {new_cell}")
            elif tag == "replace":
                doc.add_paragraph(f"Old (replaced): {old_cell}", style='Color')
                doc.add_paragraph(f"New (replaced): {new_cell}", style='Color')

    # Save the document to a file
    report_path = '/mnt/data/comparison_report.docx'
    doc.save(report_path)
    return report_path


# ------------------------
# Text extraction
# ------------------------

def extract_text_from_word(doc_file):
    doc = Document(doc_file)
    content = []

    def iter_block_items(parent):
        for child in parent.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                content.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    content.append(" | ".join(cells))
    return content


def extract_header_footer(doc_file):
    doc = Document(doc_file)
    header_text = []
    footer_text = []

    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                header_text.append(para.text.strip())
        for para in section.footer.paragraphs:
            if para.text.strip():
                footer_text.append(para.text.strip())

    return header_text, footer_text


def preprocess_image(image):
    gray_image = cv2.cvtColor(np.array(image), cv2.COLOR_BGR2GRAY)
    _, thresh_image = cv2.threshold(gray_image, 150, 255, cv2.THRESH_BINARY)
    return thresh_image


def extract_text_from_pdf(file, lang='eng', max_pages=5):
    images = convert_from_path(file, first_page=1, last_page=max_pages)
    all_text = []
    for image in images:
        processed_image = preprocess_image(image)
        text = pytesseract.image_to_string(processed_image, lang=lang)
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        all_text.extend(lines)
    return all_text


# ------------------------
# Compare Tables
# ------------------------

def extract_table_data(doc_file):
    doc = Document(doc_file)
    table_data = []

    # Loop through all tables in the document
    for table in doc.tables:
        # for row in table.rows:
        #     row_data = []
        #     for cell in row.cells:
        #         # Avoid duplicate reads by grabbing all text at once per cell
        #         cell_text = cell.text.strip()
        #         row_data.append(cell_text)
        #     table_data.append(row_data)

        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_cells)
        table_data.append(table_rows)

    return table_data


def compare_tables(old_table, new_table):
    # Compare two tables row by row, cell by cell
    table_diff = []
    changes = {"replace": 0, "insert": 0, "delete": 0}

    max_rows = max(len(old_table), len(new_table))
    for i in range(max_rows):
        old_row = old_table[i] if i < len(old_table) else []
        new_row = new_table[i] if i < len(new_table) else []

        max_cells = max(len(old_row), len(new_row))
        row_diff = []

        for j in range(max_cells):
            old_cell = old_row[j] if j < len(old_row) else ""
            new_cell = new_row[j] if j < len(new_row) else ""

            if old_cell != new_cell:
                changes["replace"] += 1
                row_diff.append((old_cell, new_cell, "replace"))
            else:
                row_diff.append((old_cell, new_cell, "equal"))

        table_diff.append(row_diff)

    return table_diff, changes


def display_table_side_by_side(old_table, new_table):
    """Display two tables side by side with highlighted differences.

    Args:
        old_table: List of lists representing the original table
        new_table: List of lists representing the modified table

    Returns:
        Dictionary of change statistics
    """
    # Add progress bar for better UX
    progress_bar = st.progress(0)

    try:
        # Compare tables and get differences
        table_diff, changes = compare_tables(old_table, new_table)

        # Debug logging (consider using logging instead of print)
        st.session_state.debug_info = {
            'old_table_len': len(old_table),
            'new_table_len': len(new_table),
            'table_diff_len': len(table_diff)
        }

        # Display tables side by side
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Old Table**")
            st.markdown("<style>.stMarkdown { margin-bottom: 0.5rem; }</style>", unsafe_allow_html=True)

        with col2:
            st.markdown("**New Table**")
            st.markdown("<style>.stMarkdown { margin-bottom: 0.5rem; }</style>", unsafe_allow_html=True)

        max_rows = max(len(old_table), len(new_table))
        max_columns = max(len(old_table[0]), len(new_table[0])) if max_rows > 0 else 0

        # Use zip_longest for safer iteration
        from itertools import zip_longest

        for i, (old_row, new_row) in enumerate(zip_longest(old_table, new_table, fillvalue=[])):
            # Update progress
            progress_bar.progress((i + 1) / max_rows)

            # Ensure rows have same number of columns
            old_row = old_row if i < len(old_table) else [""] * max_columns
            new_row = new_row if i < len(new_table) else [""] * max_columns

            # Create row header with number
            row_header = f"<div style='padding:4px; font-weight:bold; background-color:#f0f0f0'>Row {i + 1}</div>"

            # Initialize cell containers
            old_cells = [row_header]
            new_cells = [row_header]

            # Process each cell
            for j in range(max(len(old_row), len(new_row))):
                old_cell = old_row[j] if j < len(old_row) else ""
                new_cell = new_row[j] if j < len(new_row) else ""

                # Get diff status if available
                diff_status = "equal"
                if i < len(table_diff) and j < len(table_diff[i]):
                    _, _, diff_status = table_diff[i][j]

                # Apply appropriate styling
                if diff_status == "replace":
                    old_style = "padding:4px; background-color:#ffcccc"
                    new_style = "padding:4px; background-color:#cce5ff"
                else:
                    old_style = new_style = "padding:4px"

                old_cells.append(f"<div style='{old_style}'>{old_cell}</div>")
                new_cells.append(f"<div style='{new_style}'>{new_cell}</div>")

            # Display the rows
            with col1:
                st.markdown(" | ".join(old_cells), unsafe_allow_html=True)
            with col2:
                st.markdown(" | ".join(new_cells), unsafe_allow_html=True)

        # Calculate and display similarity
        total_changes = sum(changes.values())
        similarity = (1 - (changes["replace"] / total_changes)) * 100 if total_changes > 0 else 100

        progress_bar.empty()

        st.success(
            f"**Table Comparison Complete**  \n"
            f"Similarity: {similarity:.2f}%  \n"
            f"Changes: {changes['replace']} replacements"
        )

        return changes

    except Exception as e:
        progress_bar.empty()
        st.error(f"Error comparing tables: {str(e)}")
        logging.exception("Table comparison failed")
        return {"replace": 0, "insert": 0, "delete": 0}


# ------------------------
# Diff logic
# ------------------------

def get_aligned_diff(old_lines, new_lines):
    matcher = difflib.SequenceMatcher(None, old_lines, new_lines)
    aligned_diff = []
    changes = {"replace": 0, "insert": 0, "delete": 0, "equal": 0}

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        max_len = max(i2 - i1, j2 - j1)
        changes[tag] += max_len
        for i in range(max_len):
            old_line = old_lines[i1 + i] if i1 + i < i2 else ""
            new_line = new_lines[j1 + i] if j1 + i < j2 else ""
            aligned_diff.append((old_line, new_line, tag))
    return aligned_diff, changes


def highlight_entities_diff(old_line, new_line, old_entities, new_entities):
    diff = difflib.SequenceMatcher(None, old_line.split(), new_line.split())
    old_html = ""
    new_html = ""

    # To keep track of added, removed, and replaced entities
    added_entities = []
    removed_entities = []
    replaced_entities = []
    moved_entities = []

    for tag, i1, i2, j1, j2 in diff.get_opcodes():
        old_words = old_line.split()[i1:i2]
        new_words = new_line.split()[j1:j2]

        if tag == "equal":
            old_html += " " + " ".join(old_words)
            new_html += " " + " ".join(new_words)
        elif tag == "delete":
            old_html += ' <span style="background-color:#ffcccc">' + " ".join(old_words) + "</span>"
            removed_entities.append(" ".join(old_words))
        elif tag == "insert":
            new_html += ' <span style="background-color:#cce5ff">' + " ".join(new_words) + "</span>"
            added_entities.append(" ".join(new_words))
        elif tag == "replace":
            old_html += ' <span style="background-color:#ffcccc">' + " ".join(old_words) + "</span>"
            new_html += ' <span style="background-color:#cce5ff">' + " ".join(new_words) + "</span>"
            replaced_entities.append((" ".join(old_words), " ".join(new_words)))

    # Detect moved entities (entities that exist in both versions but have different positions)
    for old_entity in old_entities:
        if old_entity in new_entities and old_entity not in moved_entities:
            old_pos = old_line.find(old_entity)
            new_pos = new_line.find(old_entity)
            if old_pos != new_pos:  # Entity moved
                moved_entities.append((old_entity, old_pos, new_pos))

    return old_html.strip(), new_html.strip(), added_entities, removed_entities, replaced_entities, moved_entities


def compare_side_by_side(old_lines, new_lines, old_entities, new_entities, title="Comparison"):
    st.markdown(f"### {title}")
    aligned, stats = get_aligned_diff(old_lines, new_lines)

    if st.checkbox("Show only changed lines"):
        aligned = [item for item in aligned if item[2] != "equal"]

    col1, col2 = st.columns(2)
    col1.markdown("**Source**")
    col2.markdown("**Output**")

    # Store diff lines and their unique IDs
    diff_anchors = []
    added_entities, removed_entities, replaced_entities, moved_entities = [], [], [], []

    for idx, (old_line, new_line, _) in enumerate(aligned):
        # Generate a unique anchor for each line
        anchor_id = f"diff_{idx}"
        diff_anchors.append(anchor_id)

        old_html, new_html, added, removed, replaced, moved = highlight_entities_diff(old_line, new_line, old_entities,
                                                                                      new_entities)

        added_entities.extend(added)
        removed_entities.extend(removed)
        replaced_entities.extend(replaced)
        moved_entities.extend(moved)

        # Wrap each diff with an anchor
        col1.markdown(f"<div id='{anchor_id}' style='padding:4px'>{old_html}</div>", unsafe_allow_html=True)
        col2.markdown(f"<div id='{anchor_id}' style='padding:4px'>{new_html}</div>", unsafe_allow_html=True)

    total = sum(stats.values())
    similarity = (stats["equal"] / total * 100) if total > 0 else 100
    st.info(
        f"**Similarity:** {similarity:.2f}% — {stats['replace']} replacements, {stats['insert']} insertions, {stats['delete']} deletions")

    # Dropdown for navigating through diffs
    selected_diff = st.selectbox("🔎 Select a difference to jump to:", ["-- Select Diff --"] + diff_anchors)

    # Scroll to the selected difference if user selects one
    if selected_diff != "-- Select Diff --":
        st.markdown(f'<a href="#{selected_diff}" style="color:transparent;">Jump to Diff</a>', unsafe_allow_html=True)

    # After comparison, display the added, removed, replaced, and moved entities
    st.markdown("---")
    st.markdown("### 🧠 Named Entity Changes")

    if added_entities:
        if st.checkbox("➕ Added Entities"):
            st.markdown("**➕ Added Entities:**")
            for entity in added_entities:
                st.markdown(f"- {entity}")

    if removed_entities:
        if st.checkbox("➖ Removed Entities"):
            st.markdown("**➖ Removed Entities:**")
            for entity in removed_entities:
                st.markdown(f"- {entity}")

    if replaced_entities:
        if st.checkbox("♻️ Show Replaced Entities"):
            st.markdown("**♻️ Replaced Entities:**")
            for old_ent, new_ent in replaced_entities:
                st.markdown(f"- {old_ent} ➡️ {new_ent}")

    if moved_entities:
        if st.checkbox("🔀 Show Moved Entities"):
            st.markdown("**🔀 Moved Entities:**")
            for entity, old_pos, new_pos in moved_entities:
                st.markdown(f"- {entity} (from position {old_pos} to position {new_pos})")

    return stats, old_lines, new_lines, aligned


def compare_metadata(old_doc, new_doc):
    old_header, old_footer = extract_header_footer(old_doc)
    new_header, new_footer = extract_header_footer(new_doc)

    compare_side_by_side(old_header, new_header, title="Header Comparison")
    compare_side_by_side(old_footer, new_footer, title="Footer Comparison")


# ------------------------
# NLP Insights
# ------------------------

def get_named_entities_diff(old_lines, new_lines):
    old_text = "\n".join(old_lines)
    new_text = "\n".join(new_lines)

    old_doc = nlp(old_text)
    new_doc = nlp(new_text)

    # List of (text, label, start_char) tuples
    old_entities = [(ent.text.strip(), ent.label_, ent.start_char) for ent in old_doc.ents]
    new_entities = [(ent.text.strip(), ent.label_, ent.start_char) for ent in new_doc.ents]

    # Entity sets for matching
    old_set = {(text, label) for text, label, _ in old_entities}
    new_set = {(text, label) for text, label, _ in new_entities}

    # Detect added and removed
    added = new_set - old_set
    removed = old_set - new_set

    # Detect replaced (same label, different text)
    old_by_label = {}
    new_by_label = {}
    for text, label, _ in old_entities:
        old_by_label.setdefault(label, set()).add(text)
    for text, label, _ in new_entities:
        new_by_label.setdefault(label, set()).add(text)

    replaced = []
    for label in set(old_by_label.keys()).intersection(new_by_label.keys()):
        old_texts = old_by_label[label]
        new_texts = new_by_label[label]
        replaced_texts = (old_texts - new_texts) & (new_texts - old_texts)
        for old_text in replaced_texts:
            for new_text in replaced_texts:
                if old_text != new_text:
                    replaced.append((old_text, new_text, label))

    # Detect moved (same text/label but far apart in character position)
    moved = []
    threshold = int(len(new_text) * 0.05)  # e.g. 5% document length movement
    for ent in new_entities:
        text, label, new_pos = ent
        for old_ent in old_entities:
            if old_ent[0] == text and old_ent[1] == label:
                old_pos = old_ent[2]
                if abs(new_pos - old_pos) > threshold:
                    moved.append((text, label, old_pos, new_pos))
                break

    return {
        "added": list(added),
        "removed": list(removed),
        "replaced": replaced,
        "moved": moved
    }


def get_voice_diff(old_lines, new_lines):
    old_voice = []
    new_voice = []

    for old_line, new_line in zip(old_lines, new_lines):
        old_doc = nlp(old_line)
        new_doc = nlp(new_line)

        old_voice.append("passive" if any(tok.dep_ == "auxpass" for tok in old_doc) else "active")
        new_voice.append("passive" if any(tok.dep_ == "auxpass" for tok in new_doc) else "active")

    return old_voice, new_voice


def format_entity_changes(entities_diff):
    """Format entity changes for the AI prompt in a structured way."""
    sections = []

    if entities_diff["added"]:
        sections.append(
            "Added Entities:\n- " + "\n- ".join([f"{text} ({label})" for text, label in entities_diff["added"]]))

    if entities_diff["removed"]:
        sections.append(
            "Removed Entities:\n- " + "\n- ".join([f"{text} ({label})" for text, label in entities_diff["removed"]]))

    if entities_diff["replaced"]:
        sections.append("Replaced Entities:\n- " + "\n- ".join(
            [f"{old} → {new} ({label})" for old, new, label in entities_diff["replaced"]]))

    if entities_diff["moved"]:
        sections.append(
            "Moved Entities:\n- " + "\n- ".join([f"{text} ({label}) moved from position {old_pos} to {new_pos}"
                                                 for text, label, old_pos, new_pos in entities_diff["moved"]]))

    return "\n\n".join(sections) if sections else "No significant entity changes"


def create_summary_prompt(old_lines, new_lines, changes, entities_diff):
    """Generate a structured prompt for the AI summary."""
    return f"""
Document Comparison Report - Analysis Request:

1. Change Statistics:
- Total Changes: {sum(changes.values())}
- Replacements: {changes['replace']}
- Insertions: {changes['insert']}
- Deletions: {changes['delete']}

2. Key Entity Modifications:
{format_entity_changes(entities_diff)}

3. Content Samples:
Old Version (Excerpt):
{chr(10).join(old_lines[:10])}

New Version (Excerpt):
{chr(10).join(new_lines[:10])}

Instructions for the AI:
- Provide a concise, professional summary (3-5 sentences).
- Highlight the most impactful changes (ranked by significance).
- Mention any patterns (e.g., systematic replacements of terms).
- Note sensitive entities (names, dates, figures) that changed.
- Use bullet points for clarity if needed.
"""


# ------------------------
# AI Summary
# ------------------------

def generate_ai_summary(old_lines, new_lines, changes):
    """Placeholder for future AI integration"""
    return "🔧 AI summary feature is currently disabled. We'll implement this soon with enhanced analysis!"


# ------------------------
# Streamlit App UI
# ------------------------

st.set_page_config(page_title="Smart Document Comparator", layout="wide")
st.title("📄 AI-Enhanced Document Comparator (Draftable Style)")

dark_mode = st.checkbox("Enable Dark Mode")
if dark_mode:
    # Apply dark mode CSS styles
    st.markdown("""
    <style>
        body {
            background-color: #121212;
            color: white;
        }

        .stApp {
            background-color: #121212;
        }

        .stSidebar {
            background-color: #1d1d1d;
        }

        .stButton, .stCheckbox, .stRadio, .stSelectbox, .stTextInput, .stFileUploader {
            background-color: #333333;
            color: white;
        }

        .stButton:hover, .stCheckbox:hover, .stRadio:hover, .stSelectbox:hover, .stTextInput:hover, .stFileUploader:hover {
            background-color: #444444;
        }

        /* Change the sidebar background color */
        .css-1d391kg {
            background-color: #222222;
        }

        /* Set the font color for the title */
        .css-18e3th9 {
            color: white;
        }
    </style>
    """, unsafe_allow_html=True)

    # ===== ADD THIS NEW SIDEBAR SECTION =====
    with st.sidebar:
        st.markdown("### 🛠 Feature Roadmap")
        st.markdown("""
        <style>
            .feature-item {
                margin-bottom: 0.5rem;
                font-size: 0.9rem;
            }
        </style>
        <div class="feature-item">✓ Core comparison engine</div>
        <div class="feature-item">✓ Table/entity analysis</div>
        <div class="feature-item">◌ AI summary (Q3 2024)</div>
        """, unsafe_allow_html=True)
    # ===== END OF NEW SECTION =====

doc_type = st.radio("Choose document type", ["Word", "PDF"], horizontal=True)
ocr_lang = st.selectbox("OCR Language (PDF only)", ["eng", "spa", "fra", "deu"])
max_pages = st.slider("Pages to process (PDF only)", 1, 20, 5)

old_file = st.file_uploader("Upload OLD version", type=["docx", "pdf"])
new_file = st.file_uploader("Upload NEW version", type=["docx", "pdf"])

if old_file and new_file:
    with st.spinner("🔍 AI is analyzing your documents..."):
        try:
            if doc_type == "Word":
                old_text = extract_text_from_word(old_file)
                new_text = extract_text_from_word(new_file)
                old_entities = extract_entities(" ".join(old_text))
                new_entities = extract_entities(" ".join(new_text))
                old_tables = extract_table_data(old_file)
                new_tables = extract_table_data(new_file)

                # ====== ADD THE TABBED UI HERE ======
                tab1, tab2, tab3 = st.tabs(["📝 Content Comparison", "🧠 Entity Analysis", "🤖 AI Summary"])

                with tab1:
                    # Content & Table Comparison
                    stats, o_lines, n_lines, aligned = compare_side_by_side(
                        old_text, new_text, old_entities, new_entities, title="Text Comparison"
                    )

                    # Compare each table
                    for old_table, new_table in zip(old_tables, new_tables):
                        display_table_side_by_side(old_table, new_table)

                with tab2:
                    # Named Entity Changes
                    st.markdown("### 🧠 Named Entity Changes")
                    entities_diff = get_named_entities_diff(o_lines, n_lines)

                    if st.checkbox("➕ Added Entities", key="added_ents"):
                        st.write("**Added:**", [e[0] for e in entities_diff["added"]])

                    if st.checkbox("➖ Removed Entities", key="removed_ents"):
                        st.write("**Removed:**", [e[0] for e in entities_diff["removed"]])

                    if st.checkbox("♻️ Replaced Entities", key="replaced_ents"):
                        st.write("**Replaced:**", [f"{old} → {new}" for old, new, _ in entities_diff["replaced"]])

                with tab3:
                    st.markdown("### 🤖 AI Summary (Coming Soon)")
                    st.info("""
                    We're working on an advanced AI summary feature that will:
                    - Highlight key changes automatically
                    - Explain impacts in business terms
                    - Suggest action items

                    **Try these current features:**
                    - Text differences (left tab)
                    - Table comparisons (below)
                    - Entity tracking (previous tab)
                    """)

                    # Show basic stats from the comparison
                    st.metric("Total Changes", sum(stats.values()))
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Replacements", stats['replace'])
                    col2.metric("Additions", stats['insert'])
                    col3.metric("Deletions", stats['delete'])
                # ====== END OF TABBED UI ======

            else:  # PDF handling
                old_text = extract_text_from_pdf(old_file, lang=ocr_lang, max_pages=max_pages)
                new_text = extract_text_from_pdf(new_file, lang=ocr_lang, max_pages=max_pages)
                # ... (rest of your PDF comparison logic)

        except Exception as e:
            st.error(f"Error: {str(e)}")
else:
    st.info("📥 Please upload both documents to begin comparison.")
