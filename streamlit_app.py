import streamlit as st
import difflib
from docx import Document
from pdf2image import convert_from_path
import pytesseract
import openai
import os
import cv2
import numpy as np
from dotenv import load_dotenv  # 🔹 NEW

# Load API key from .env
load_dotenv()  # 🔹 NEW
openai.api_key = os.getenv("OPENAI_API_KEY")  # 🔹 NEW


# ------------------------
# TEXT & TABLE EXTRACTION
# ------------------------

def extract_text_from_word(doc_file):
    doc = Document(doc_file)
    content = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append(text)
    return content


def extract_tables_from_word(doc_file):
    doc = Document(doc_file)
    tables = []

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return tables


def preprocess_image(image):
    gray_image = cv2.cvtColor(np.array(image), cv2.COLOR_BGR2GRAY)
    _, thresh_image = cv2.threshold(gray_image, 150, 255, cv2.THRESH_BINARY)
    return thresh_image


def extract_text_from_pdf(file, lang='eng', max_pages=5):
    images = convert_from_path(file, first_page=1, last_page=max_pages)
    all_text = []
    for image in images:
        processed_image = preprocess_image(image)
        text = pytesseract.image_to_string(processed_image, lang=lang)
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        all_text.extend(lines)
    return all_text


# ------------------------
# DIFF LOGIC FOR TEXT & TABLE
# ------------------------

def get_aligned_text_diff(old_lines, new_lines):
    matcher = difflib.SequenceMatcher(None, old_lines, new_lines)
    aligned_diff = []
    stats = {"replace": 0, "insert": 0, "delete": 0, "equal": 0}

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        for i in range(max(i2 - i1, j2 - j1)):
            old = old_lines[i1 + i] if i1 + i < i2 else ""
            new = new_lines[j1 + i] if j1 + i < j2 else ""
            aligned_diff.append((old, new, tag))
            stats[tag] += 1

    return aligned_diff, stats


def get_aligned_table_diff(old_table, new_table):
    aligned_diff = []
    changes = {"replace": 0, "insert": 0, "delete": 0, "equal": 0}

    max_rows = max(len(old_table), len(new_table))
    for i in range(max_rows):
        old_row = old_table[i] if i < len(old_table) else []
        new_row = new_table[i] if i < len(new_table) else []

        max_cols = max(len(old_row), len(new_row))
        row_diff = []

        for j in range(max_cols):
            old_cell = old_row[j] if j < len(old_row) else ""
            new_cell = new_row[j] if j < len(new_row) else ""
            tag = "equal" if old_cell == new_cell else "replace"
            if old_cell == "" and new_cell:
                tag = "insert"
            elif new_cell == "" and old_cell:
                tag = "delete"
            row_diff.append((old_cell, new_cell, tag))
            changes[tag] += 1

        aligned_diff.append(row_diff)

    return aligned_diff, changes


# ------------------------
# AI SUMMARY
# ------------------------

def generate_ai_summary(old_lines, new_lines, changes):
    prompt = f"""
You're a document comparison assistant. Here's a summary of the comparison:

Changes Detected:
- {changes['replace']} replacements
- {changes['insert']} insertions
- {changes['delete']} deletions

Examples of old content:
{chr(10).join(old_lines[:5])}

Examples of new content:
{chr(10).join(new_lines[:5])}

Provide a professional and simple summary of the key differences.
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You summarize key differences in legal or technical documents."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
            max_tokens=500
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"⚠️ AI summary failed: {e}"


# ------------------------
# STREAMLIT APP UI
# ------------------------

st.set_page_config(page_title="Smart Document Comparator", layout="wide")
st.title("📄 Smart Document Comparator with AI")

doc_type = st.radio("Choose document type", ["Word", "PDF"], horizontal=True)
ocr_lang = st.selectbox("OCR Language (PDF only)", ["eng", "spa", "fra", "deu"])
max_pages = st.slider("Pages to process (PDF only)", 1, 20, 5)

old_file = st.file_uploader("Upload OLD version", type=["docx", "pdf"])
new_file = st.file_uploader("Upload NEW version", type=["docx", "pdf"])

if old_file and new_file:
    with st.spinner("🔍 Analyzing your documents..."):
        try:
            if doc_type == "Word":
                old_text = extract_text_from_word(old_file)
                new_text = extract_text_from_word(new_file)
                old_tables = extract_tables_from_word(old_file)
                new_tables = extract_tables_from_word(new_file)
            else:
                old_text = extract_text_from_pdf(old_file, lang=ocr_lang, max_pages=max_pages)
                new_text = extract_text_from_pdf(new_file, lang=ocr_lang, max_pages=max_pages)
                old_tables, new_tables = [], []

            # TEXT DIFF
            aligned_text, stats_text = get_aligned_text_diff(old_text, new_text)

            st.markdown("### 📝 Text Changes (Side by Side)")
            col1, col2 = st.columns(2)
            col1.markdown("**Old Version**")
            col2.markdown("**New Version**")

            for old_line, new_line, tag in aligned_text:
                if tag != "equal":
                    col1.markdown(f"<div style='background-color:#ffdddd;padding:6px'>{old_line}</div>",
                                  unsafe_allow_html=True)
                    col2.markdown(f"<div style='background-color:#ddffdd;padding:6px'>{new_line}</div>",
                                  unsafe_allow_html=True)

            # TABLE DIFF
            if old_tables or new_tables:
                st.markdown("### 📊 Table Comparison")

                for idx in range(min(len(old_tables), len(new_tables))):
                    st.markdown(f"#### Table {idx + 1}")
                    table_diff, _ = get_aligned_table_diff(old_tables[idx], new_tables[idx])

                    col1, col2 = st.columns(2)
                    col1.markdown("**Old Table**")
                    col2.markdown("**New Table**")

                    for row in table_diff:
                        old_row_cells = []
                        new_row_cells = []

                        for cell in row:
                            if len(cell) == 3:
                                old_val, new_val, tag = cell
                            else:
                                old_val, new_val, tag = "", "", "equal"

                            if tag == "replace":
                                old_val = f"<span style='background-color:#ffcccc'>{old_val}</span>"
                                new_val = f"<span style='background-color:#cce5ff'>{new_val}</span>"
                            elif tag == "insert":
                                new_val = f"<span style='background-color:#cce5ff'>{new_val}</span>"
                            elif tag == "delete":
                                old_val = f"<span style='background-color:#ffcccc'>{old_val}</span>"

                            old_row_cells.append(f"<td>{old_val}</td>")
                            new_row_cells.append(f"<td>{new_val}</td>")

                        col1.markdown(f"<table><tr>{''.join(old_row_cells)}</tr></table>", unsafe_allow_html=True)
                        col2.markdown(f"<table><tr>{''.join(new_row_cells)}</tr></table>", unsafe_allow_html=True)

            # AI SUMMARY
            st.markdown("---")
            st.markdown("### 🤖 AI Summary of Changes")
            ai_summary = generate_ai_summary(old_text, new_text, stats_text)
            st.success(ai_summary)

            st.success("✅ Document comparison complete!")

        except Exception as e:
            st.error(f"❌ Something went wrong: {e}")
else:
    st.info("📥 Please upload both documents to begin comparison.")
